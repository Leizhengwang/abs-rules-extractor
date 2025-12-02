import os
import io
import uuid
import zipfile
import time
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import hashlib
import json

# Try to import psutil for detailed health checks, but make it optional
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# Track app start time for uptime monitoring
APP_START_TIME = time.time()

# Ensure upload and output directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_red_pdf_contents(pdf_data, original_filename=None):
    """Extract red content from PDF and return processed PDF data"""
    doc = fitz.open(stream=pdf_data, filetype="pdf")
    
    # Configuration
    dpi_resol = 380
    red_blocks_found = 0
    total_blocks = 0
    
    # Create new PDF for red content
    new_pdf = fitz.open()
    
    print(f"Processing PDF with {doc.page_count} pages...")
    
    for page_num in range(doc.page_count):
        page_has_redtable = False
        tableposted_flag = False
        new_page = new_pdf.new_page(width=doc[page_num].rect.width, height=doc[page_num].rect.height)
        page_flag = False
        text_blocks = doc[page_num].get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
        page_red_blocks = 0
        
        print(f"Page {page_num + 1}: Found {len(text_blocks)} text blocks")
        
        # Extract red content
        for i, block in enumerate(text_blocks):
            if 'lines' not in block:
                continue
                
            total_blocks += 1
            x0, y0, x1, y1 = block['bbox']
            rect = fitz.Rect(x0, y0, x1, y1)
            block_flag = False
            block_printed = False
            block_text = ""
            
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if not text:
                        continue
                        
                    block_text += text + " "
                    font_size = span.get("size", None)
                    font_name = span.get("font", "")
                    color = fitz.sRGB_to_rgb(span["color"])
                    is_bold = "Bold" in font_name or "Black" in font_name
                    
                    # Debug: Print color information for first few blocks
                    if total_blocks <= 5:
                        print(f"  Block {total_blocks}: Text='{text[:20]}...', Color={color}, Font={font_name}, Size={font_size}")
                    
                    # Check for red color text - be more flexible with red detection
                    # Standard red: (218, 31, 51), but also check for similar reds
                    red_threshold = 50  # Allow some variation
                    is_red = (color[0] > 150 and color[1] < 100 and color[2] < 100) or color == (218, 31, 51)
                    
                    if is_red:
                        print(f"Found RED text on page {page_num + 1}: '{text}' with color {color}")
                        if text not in ["•", "●", "∙", "–"]:
                            if not page_flag:
                                print(f"Processing page number: {page_num + 1}")
                                page_flag = True
                            block_flag = True
                            page_red_blocks += 1
                            
                            # Handle tables with red text
                            if ((font_size == 9 and font_name == "TimesNewRomanPSMT") or 
                                (font_size == 12 and font_name == "Arial-ItalicMT" and text in ["2024", "2025"])) and not tableposted_flag:
                                
                                try:
                                    with pdfplumber.open(BytesIO(pdf_data)) as pdf_plumber:
                                        for page_number, page in enumerate(pdf_plumber.pages, start=1):
                                            if page_number == page_num + 1:
                                                tables = page.find_tables()
                                                if tables:
                                                    page_has_redtable = True
                                                    print(f"Found {len(tables)} tables on page {page_num + 1}")
                                                    
                                                    for table_index, table in enumerate(tables):
                                                        x0_t, top, x1_t, bottom = table.bbox
                                                        pix = doc[page_num].get_pixmap(dpi=380, clip=(x0_t, top, x1_t, bottom))
                                                        new_page.insert_image(fitz.Rect(x0_t, top, x1_t, bottom), pixmap=pix)
                                                    
                                                    tableposted_flag = True
                                except Exception as e:
                                    print(f"Error processing tables: {e}")
                    
                    elif color == (0, 0, 0) and not block_printed:
                        if is_bold and ((font_size == 14 and font_name == "Arial-Black") or 
                                       (font_size == 36 and font_name == "Arial-Black")):
                            try:
                                image = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect))
                                new_page.insert_image(fitz.Rect(rect), pixmap=image)
                                block_printed = True
                                print(f"Added bold black text block on page {page_num + 1}")
                            except Exception as e:
                                print(f"Error adding bold text: {e}")
            
            # Add block with red content
            if block_flag:
                try:
                    image = doc[page_num].get_pixmap(dpi=dpi_resol, clip=fitz.Rect(rect))
                    new_page.insert_image(fitz.Rect(rect), pixmap=image)
                    red_blocks_found += 1
                    print(f"Added red text block: '{block_text[:50]}...'")
                except Exception as e:
                    print(f"Error adding red block: {e}")
        
        print(f"Page {page_num + 1}: Added {page_red_blocks} red blocks")
    
    print(f"Total red blocks found and processed: {red_blocks_found} out of {total_blocks} total blocks")
    
    # Save to memory buffer
    buffer = BytesIO()
    new_pdf.save(buffer, garbage=4, deflate=True)
    pdf_data_result = buffer.getvalue()
    buffer.close()
    new_pdf.close()
    doc.close()
    
    return pdf_data_result

def extract_images_with_positions(pdf_data):
    """Extract images from PDF with their positions"""
    pdf_document = fitz.open(stream=pdf_data, filetype="pdf")
    images = []
    
    print(f"Extracting images from PDF with {len(pdf_document)} pages...")
    
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        image_list = page.get_images(full=True)
        
        print(f"Page {page_num + 1}: Found {len(image_list)} images")
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            image_rects = page.get_image_rects(xref)
            for rect in image_rects:
                images.append({
                    "image_bytes": image_bytes,
                    "image_ext": image_ext,
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1,
                    "width": rect.width,
                    "height": rect.height
                })
                print(f"  Image {len(images)}: {rect.width}x{rect.height} at ({rect.x0}, {rect.y0})")
    
    print(f"Total images extracted: {len(images)}")
    pdf_document.close()
    return images

def create_word_document_with_positioned_images(images, output_docx_path):
    """Create Word document with positioned images"""
    doc = Document()
    
    print(f"Creating Word document with {len(images)} images...")
    
    if len(images) == 0:
        print("WARNING: No images found to add to Word document!")
        # Add a message to the document if no images found
        doc.add_paragraph("No red text content was found in the PDF file.")
        doc.add_paragraph("This could mean:")
        doc.add_paragraph("• The PDF doesn't contain red text")
        doc.add_paragraph("• The red color doesn't match the expected RGB values")
        doc.add_paragraph("• The text is embedded as images rather than text")
    else:
        for i, image_info in enumerate(images):
            try:
                image_bytes = image_info["image_bytes"]
                x0 = image_info["x0"]
                width = image_info["width"]
                height = image_info["height"]
                
                print(f"Adding image {i+1}: {width}x{height} at position {x0}")
                
                # Add paragraph for the image
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                
                # Add image to paragraph
                image_stream = BytesIO(image_bytes)
                # Limit image width to reasonable size (max 6 inches)
                max_width = min(width / 86, 6.0)
                run.add_picture(image_stream, width=Inches(max_width))
                
                # Set horizontal position
                paragraph.alignment = 0
                paragraph.paragraph_format.left_indent = Inches(x0 / 86)
                
            except Exception as e:
                print(f"Error adding image {i+1}: {e}")
    
    doc.save(output_docx_path)
    print(f"Word document saved: {output_docx_path}")
    return output_docx_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        return jsonify({'error': 'No files selected'}), 400
    
    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({'error': 'No files selected'}), 400
    
    processed_files = []
    job_id = str(uuid.uuid4())
    
    try:
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                base_name = filename.rsplit('.', 1)[0]
                
                # Read PDF data
                pdf_data = file.read()
                
                # Extract red content
                processed_pdf_data = extract_red_pdf_contents(pdf_data, filename)
                
                # Extract images from processed PDF
                images = extract_images_with_positions(processed_pdf_data)
                
                # Create Word document
                word_filename = f"word_output_{base_name}.docx"
                # Limit filename length to avoid path issues
                if len(word_filename) > 50:
                    word_filename = f"word_output_{base_name[:20]}.docx"
                
                word_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{job_id}_{word_filename}")
                print(f"Creating Word document at: {word_path}")
                
                create_word_document_with_positioned_images(images, word_path)
                print(f"Word document created successfully: {os.path.exists(word_path)}")
                
                processed_files.append({
                    'original_name': filename,
                    'word_file': f"{job_id}_{word_filename}",
                    'status': 'completed'
                })
        
        return jsonify({
            'success': True,
            'job_id': job_id,
            'files': processed_files,
            'message': f'Successfully processed {len(processed_files)} files'
        })
    
    except Exception as e:
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/test_download')
def test_download():
    """Test route to check if downloads work"""
    try:
        output_files = os.listdir(app.config['OUTPUT_FOLDER'])
        if output_files:
            test_file = output_files[0]
            file_path = os.path.join(app.config['OUTPUT_FOLDER'], test_file)
            return send_file(file_path, as_attachment=True, download_name=test_file)
        else:
            return "No files in output folder"
    except Exception as e:
        return f"Error: {str(e)}"

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        print(f"Download request for: {filename}")
        print(f"Looking for file at: {file_path}")
        print(f"File exists: {os.path.exists(file_path)}")
        
        if os.path.exists(file_path):
            print(f"File size: {os.path.getsize(file_path)} bytes")
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            print(f"File not found! Files in output folder:")
            if os.path.exists(app.config['OUTPUT_FOLDER']):
                for f in os.listdir(app.config['OUTPUT_FOLDER']):
                    print(f"  - {f}")
            else:
                print(f"Output folder doesn't exist: {app.config['OUTPUT_FOLDER']}")
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        print(f"Download error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_all/<job_id>')
def download_all(job_id):
    try:
        # Create zip file with all processed documents
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename in os.listdir(app.config['OUTPUT_FOLDER']):
                if filename.startswith(job_id):
                    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
                    # Remove job_id prefix from filename in zip
                    archive_name = filename.replace(f"{job_id}_", "")
                    zip_file.write(file_path, archive_name)
        
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'extracted_documents_{job_id}.zip'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/status/<job_id>')
def get_status(job_id):
    # Check status of processing job
    files_in_output = [f for f in os.listdir(app.config['OUTPUT_FOLDER']) if f.startswith(job_id)]
    return jsonify({
        'job_id': job_id,
        'status': 'completed' if files_in_output else 'processing',
        'files_count': len(files_in_output)
    })

# ============================================================================
# HEALTH MONITORING ENDPOINTS
# ============================================================================

@app.route('/health')
def health_check():
    """Basic health check endpoint for Azure health checks"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.utcnow().isoformat(),
        'uptime_seconds': int(time.time() - APP_START_TIME),
        'service': 'ABS Rules Red Text Extractor'
    }), 200

@app.route('/health/liveness')
def liveness_check():
    """Kubernetes-style liveness probe - is the app alive?"""
    return jsonify({
        'status': 'alive',
        'timestamp': datetime.utcnow().isoformat()
    }), 200

@app.route('/health/readiness')
def readiness_check():
    """Kubernetes-style readiness probe - is the app ready to serve requests?"""
    try:
        # Check if required directories exist and are writable
        upload_ready = os.path.exists(app.config['UPLOAD_FOLDER']) and os.access(app.config['UPLOAD_FOLDER'], os.W_OK)
        output_ready = os.path.exists(app.config['OUTPUT_FOLDER']) and os.access(app.config['OUTPUT_FOLDER'], os.W_OK)
        
        if upload_ready and output_ready:
            return jsonify({
                'status': 'ready',
                'timestamp': datetime.utcnow().isoformat(),
                'upload_folder': app.config['UPLOAD_FOLDER'],
                'output_folder': app.config['OUTPUT_FOLDER']
            }), 200
        else:
            return jsonify({
                'status': 'not_ready',
                'upload_folder_ready': upload_ready,
                'output_folder_ready': output_ready,
                'timestamp': datetime.utcnow().isoformat()
            }), 503
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e),
            'timestamp': datetime.utcnow().isoformat()
        }), 503

@app.route('/health/detailed')
def detailed_health_check():
    """Detailed health check with system metrics (if psutil is available)"""
    try:
        health_data = {
            'status': 'healthy',
            'timestamp': datetime.utcnow().isoformat(),
            'uptime_seconds': int(time.time() - APP_START_TIME),
            'service': 'ABS Rules Red Text Extractor',
            'version': '2.0',
            'application': {
                'upload_folder': app.config['UPLOAD_FOLDER'],
                'upload_folder_exists': os.path.exists(app.config['UPLOAD_FOLDER']),
                'output_folder': app.config['OUTPUT_FOLDER'],
                'output_folder_exists': os.path.exists(app.config['OUTPUT_FOLDER']),
                'max_content_length_mb': app.config['MAX_CONTENT_LENGTH'] / (1024 * 1024)
            }
        }
        
        # Add system metrics if psutil is available
        if PSUTIL_AVAILABLE:
            try:
                cpu_percent = psutil.cpu_percent(interval=0.1)
                memory = psutil.virtual_memory()
                disk = psutil.disk_usage('/')
                
                health_data['system'] = {
                    'cpu_percent': round(cpu_percent, 2),
                    'memory_percent': round(memory.percent, 2),
                    'memory_available_mb': round(memory.available / (1024 * 1024), 2),
                    'memory_total_mb': round(memory.total / (1024 * 1024), 2),
                    'disk_percent': round(disk.percent, 2),
                    'disk_free_gb': round(disk.free / (1024 * 1024 * 1024), 2),
                    'disk_total_gb': round(disk.total / (1024 * 1024 * 1024), 2)
                }
                
                # Determine if system is healthy based on thresholds
                is_healthy = (
                    cpu_percent < 90 and 
                    memory.percent < 90 and 
                    disk.percent < 90 and
                    health_data['application']['upload_folder_exists'] and 
                    health_data['application']['output_folder_exists']
                )
                
                health_data['status'] = 'healthy' if is_healthy else 'degraded'
                status_code = 200 if is_healthy else 503
                
            except Exception as e:
                health_data['system'] = {'error': f'Failed to get system metrics: {str(e)}'}
                status_code = 200  # Still return 200 for app health even if system metrics fail
        else:
            health_data['system'] = {
                'note': 'Install psutil for detailed system metrics: pip install psutil'
            }
            status_code = 200
        
        return jsonify(health_data), status_code
        
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.utcnow().isoformat()
        }), 503

@app.route('/metrics')
def metrics():
    """Prometheus-style metrics endpoint"""
    try:
        # Count files in upload and output folders
        upload_files = len([f for f in os.listdir(app.config['UPLOAD_FOLDER']) if os.path.isfile(os.path.join(app.config['UPLOAD_FOLDER'], f))])
        output_files = len([f for f in os.listdir(app.config['OUTPUT_FOLDER']) if os.path.isfile(os.path.join(app.config['OUTPUT_FOLDER'], f))])
        
        uptime = int(time.time() - APP_START_TIME)
        
        metrics_text = f"""# HELP app_uptime_seconds Application uptime in seconds
# TYPE app_uptime_seconds counter
app_uptime_seconds {uptime}

# HELP app_upload_files_total Total number of files in upload folder
# TYPE app_upload_files_total gauge
app_upload_files_total {upload_files}

# HELP app_output_files_total Total number of files in output folder
# TYPE app_output_files_total gauge
app_output_files_total {output_files}
"""
        
        if PSUTIL_AVAILABLE:
            cpu_percent = psutil.cpu_percent(interval=0.1)
            memory = psutil.virtual_memory()
            disk = psutil.disk_usage('/')
            
            metrics_text += f"""
# HELP system_cpu_percent CPU usage percentage
# TYPE system_cpu_percent gauge
system_cpu_percent {cpu_percent}

# HELP system_memory_percent Memory usage percentage
# TYPE system_memory_percent gauge
system_memory_percent {memory.percent}

# HELP system_memory_available_bytes Available memory in bytes
# TYPE system_memory_available_bytes gauge
system_memory_available_bytes {memory.available}

# HELP system_disk_percent Disk usage percentage
# TYPE system_disk_percent gauge
system_disk_percent {disk.percent}

# HELP system_disk_free_bytes Free disk space in bytes
# TYPE system_disk_free_bytes gauge
system_disk_free_bytes {disk.free}
"""
        
        return metrics_text, 200, {'Content-Type': 'text/plain; charset=utf-8'}
        
    except Exception as e:
        return f"# Error generating metrics: {str(e)}", 500, {'Content-Type': 'text/plain; charset=utf-8'}

# ============================================================================
# END HEALTH MONITORING ENDPOINTS
# ============================================================================

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8000)