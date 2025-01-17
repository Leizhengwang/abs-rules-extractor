import io
import os
import pymupdf
import json
import matplotlib
import numpy as np
import fitz  
from docx import Document
from PIL import Image
from io import BytesIO 
from docx.shared import Inches
import cv2
matplotlib.use('TkAgg',force=True)

# 1. use pymupdf to detect every red content, filter some noise and fake dot red content
# 2. from the font size, bold and name know where the rec content should be, extract level 1, 2, 3 flag , and figure information and check they are red colored or not, save it if not red
# 3. when encounting every red print out content, attach level 1 , 2, 3 and figure if necssary ( if they not covered by red content)
# 4. firstly print out all extracted red contents in a temporary pdf, finally stack all these net pdf contents into a word file


# 1. load pdf file and extract all figures
pdf = "1-mvr-part-4-jul24_432.pdf"
doc = fitz.open(pdf)

#firstly extract all figures and save in a figure folder
figure_dpi = {}
def extract_figures_from_pdf(pdf_path, output_folder):
    doc = fitz.open(pdf_path)
    # Create an output folder for saving images
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    image_counter = 0 # Figure number
    # Iterate through each page of the PDF
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)  # Load the page
        # Extract images (figures) from the page
        image_list = page.get_images(full=True)

        if len(image_list) <= 1: # every pafge has a page format figure which is not a real figure
            continue
        for img_index, img in enumerate(image_list[:-1]):
            xref = img[0]  # Get image reference
            base_image = doc.extract_image(xref)  # Extract image as a dictionary
            # Get image bytes and convert to PIL Image
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            width, height = image.size
            dpi = image.info.get("dpi", (96.012, 96.012))
            rect = page.get_image_bbox(img)   ## save the image and locatio in this page
            rect_data = {
                        "x0": rect.x0,
                        "y0": rect.y0,
                        "x1": rect.x1,
                        "y1": rect.y1
                        }
            rect_filename = os.path.join(output_folder, f"FIGURE {image_counter+1}.json")    
            with open(rect_filename, 'w') as f:
                json.dump(rect_data, f)
            # Save the image as PNG
            image_filename = os.path.join(output_folder, f"FIGURE {image_counter+1}.png")
            image.save(image_filename, dpi = (96.012, 96.012))
            image_counter += 1
            print(f"image {image_counter}: {width} * {height} pixels, DPI:{dpi} {dpi[0]}")
            figure_dpi[image_counter] = dpi[0]
    
pdf_path = pdf  # Replace with your PDF file path
output_folder = 'extracted_images'  # Folder to store extracted images
extract_figures_from_pdf(pdf_path, output_folder)


def deblur_image(image_bytes):
    # Convert bytes to OpenCV image (numpy array)
    image = Image.open(BytesIO(image_bytes))
    image = np.array(image)
    
    # Example kernel for sharpening/deblurring
    kernel = np.array([[0, -1, 0], [-1, 5,-1], [0, -1, 0]])  # Simple sharpening kernel

    # Apply the kernel to the image
    deblurred_image = cv2.filter2D(image, -1, kernel)
    
    return deblurred_image

font_size_flag = False



# 2, loop page by page , detect all red contents
new_pdf = fitz.open()
for page_num in range(doc.page_count):  #doc.page_count
    new_page =new_pdf.new_page(width = doc[page_num].rect.width, height =doc[page_num].rect.height)
    page_flag = False
    text_blocks = doc[page_num].get_text("dict", flags=pymupdf.TEXTFLAGS_TEXT)["blocks"]
    flag_2 = False   # label level 2 outline is black or red
    flag_3 = False   # label level 3 outline is black or red

    

    # extract red content
    for i, block in enumerate(text_blocks):
        x0, y0, x1, y1, = block['bbox']
        rect = fitz.Rect(x0, y0, x1, y1)
        level_font_dict = set()
        block_flag = False
        flag_2_red_done = False 
        flag_3_red_done = False
        figure_flag = False
        figure_list = []

        

        for line in block["lines"]:
            for span in line["spans"]:
                text = span["text"]
                font_size = span.get("size", None)
                font_name = span.get("font")
                color = pymupdf.sRGB_to_rgb(span["color"])
                is_bold = "Bold" in font_name or"Black" in font_name
                level_font_dict.add((font_size, font_name))
                # red color text or not
                if color == (218, 31, 51):
                    if text == "•" or text == "●":
                        continue
                    else:
                        if not page_flag:
                            print(f"page numer: {page_num+1}")
                            page_flag = True
                        block_flag = True

        if block_flag:  # detect this red content itself is level one or level two outline or not
            if (11.0, "Arial-BoldMT") in level_font_dict: 
                flag_2_red_done = True   # level 2 outline labled by red or not
            elif (10.0, "Arial-BoldMT") in level_font_dict: 
                flag_3_red_done = True


        ### write down all font info from this block, extract level 2 and level 3 stuff, if if already labelled by red, then start from 0
        level_font_dict_all = set()
    
        for line in block["lines"]:
            for span in line["spans"]:
                text = span["text"]
                font_size = span.get("size", None)
                font_name = span.get("font")
                color = pymupdf.sRGB_to_rgb(span["color"])
                is_bold = "Bold" in font_name or "Black" in font_name
                level_font_dict_all.add((font_size, font_name))

                if font_size == 10:
                    font_size_flag = True

        if block_flag:
            if (11.0, "Arial-BoldMT") in level_font_dict_all:   
                flag_2 , flag_3 = False , False
                font_size_flag = False

            if (10.0, "Arial-BoldMT") in level_font_dict_all:  
                flag_3 = False 
            if (12.0, "Arial-BoldMT") in level_font_dict_all:  # detect is figure related text or not,
                for line in block["lines"][:1]:
                    for span in line["spans"]:
                        text = span["text"]
                        if text[:6] =="FIGURE":
                            figure_list.append(text)
                figure_flag = True

        dpi_resol = 580

        #if goes to another chapter, then all level 2 and level 3 flag == 0, form the beginning set up 0
        if (12.0, "Arial-Black") in level_font_dict_all:
            flag_2 , flag_3 = False , False
        else:       # save the level 2 and level 3 outline if not in red
            if (11.0, "Arial-BoldMT") in level_font_dict_all and not flag_2_red_done:
                rect_2 = rect
                image_2 = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect_2))
                image_2.save("output_image2.png")
                rect_data = {
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1
                    }
                with open('rect_for_crop_2.json', 'w') as f:
                    json.dump(rect_data, f)
                flag_2 = True

            if (10.0, "Arial-BoldMT") in level_font_dict_all and not flag_3_red_done:
                rect_3 = rect
                image_3 = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect_3))
                image_3.save("output_image3.png")
                rect_data = {
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1
                    } 
                with open('rect_for_crop_3.json', 'w') as f:
                    json.dump(rect_data, f)
                #new_page.draw_rect(rect_3, color = (1, 0, 0), width =1)
                flag_3 = True

        # finally decide what exactly need to print out for this red content area in total
        if block_flag:
            if not flag_2_red_done and flag_2 and os.path.exists('rect_for_crop_2.json'): # if level 2 not red color, print out it
                with open('rect_for_crop_2.json', 'r') as f:
                    loaded_rect_data = json.load(f)  
                rect_2 = fitz.Rect(loaded_rect_data['x0'], loaded_rect_data['y0'], loaded_rect_data['x1'], loaded_rect_data['y1'])    
                new_page.insert_image(fitz.Rect(rect_2), pixmap = fitz.Pixmap("output_image2.png"))
                os.remove('rect_for_crop_2.json')
                
            if not flag_3_red_done and flag_3 and os.path.exists('rect_for_crop_3.json'): # if level 3 not red color, print out it
                with open('rect_for_crop_3.json', 'r') as f:
                    loaded_rect_data = json.load(f)  
                rect_3 = fitz.Rect(loaded_rect_data['x0'], loaded_rect_data['y0'], loaded_rect_data['x1'], loaded_rect_data['y1'])
                new_page.insert_image(fitz.Rect(rect_3), pixmap = fitz.Pixmap("output_image3.png"))
                os.remove('rect_for_crop_3.json')
            
            if figure_flag: # if figure title is red color , print out this figure
                for figure_name in figure_list:
                    if figure_name[-1] == " ":
                        figure_name = figure_name[:-1]
                    rect_filename = os.path.join(output_folder, figure_name + ".json") 
                    image_filename = os.path.join(output_folder, figure_name +".png")
                    with open(rect_filename, 'r') as f:
                        loaded_rect_data = json.load(f)                   
                    rect4 = fitz.Rect(loaded_rect_data['x0'], loaded_rect_data['y0'], loaded_rect_data['x1'], loaded_rect_data['y1'])
                    new_page.insert_image(fitz.Rect(rect4), pixmap = fitz.Pixmap(image_filename))   

            image = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect))
            new_page.insert_image(fitz.Rect(rect), pixmap = image)
            
    new_pdf_path = "New_red_contents_v3.pdf"
    new_pdf.save(new_pdf_path)


# 3. scan the generated pdf intermediate file and stack the contents into word format
#  firstly scan all as images, need to scale down for some large images which is too large for the word page
def extract_images_from_pdf(pdf_path):
    # Open the PDF
    doc = fitz.open(pdf_path)

    images = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            xref = img[0]  # Image xref
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            images.append(image_bytes)  
    return images

pdf_path = 'New_red_contents_v3.pdf'
images = extract_images_from_pdf(pdf_path)


# remove intermediate files and folders by using  os.rmdir
if os.path.exists('rect_for_crop_2.json'): 
    os.remove('rect_for_crop_2.json')

if os.path.exists('rect_for_crop_3.json'): 
    os.remove('rect_for_crop_3.json')

if os.path.exists('output_image2.png'): 
    os.remove('output_image2.png')

if os.path.exists('output_image3.png'): 
    os.remove('output_image3.png')

#if os.path.exists('New_red_contents_v3.pdf'): 
    #os.remove('New_red_contents_v3.pdf')


#  print pdf red contects to word
scale_ratio = 1.05   # scaling factor is the image is too large

# Function to insert images into a Word document
def insert_images_into_word(image_data, word_output_path):
    doc = Document()
    
    for img_data in image_data:
        # Open image from byte data
        img = Image.open(io.BytesIO(img_data))
        
        # Save image to temporary file for insertion
        temp_image_path = "temp_image.jpg"
        img.save(temp_image_path)

        dpi_original = img.info.get("dpi", (72, 72))
        print(dpi_original)
        #if dpi_original[0] == 96.012:
            # Add image to Word document with original size
            #doc.add_paragraph()  # Add a paragraph for spacing
            #doc.add_picture(temp_image_path, width=Inches(scale_ratio * img.width / dpi_resol), height=Inches(scale_ratio * img.height / dpi_resol))  # Using 580 DPI for scaling
        #else:
        doc.add_picture(temp_image_path, width=Inches(scale_ratio * img.width /dpi_resol), height=Inches(scale_ratio * img.height / dpi_resol))  # Using real figure DPI for scaling
        # Optionally, you can add a caption or other elements if needed
    
    # Save the Word document
    doc.save(word_output_path)
    print(f"Word document saved at: {word_output_path}")


# Main function to extract images and create Word document
def main(pdf_path, word_output_path):
    # Step 1: Extract images from the PDF
    image_data = extract_images_from_pdf(pdf_path)
    
    # Step 2: Insert images into the Word document
    insert_images_into_word(image_data, word_output_path)


# Example usage
pdf_path = 'New_red_contents_v3.pdf'  # Provide your PDF file path
word_output_path = 'output.docx'  # Specify the output Word file path
main(pdf_path, word_output_path)
