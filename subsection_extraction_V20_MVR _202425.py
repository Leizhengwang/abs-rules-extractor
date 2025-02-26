#major changes for different version
# 1, this rule table title  12.0, Arial-ItalicMT   12.0, Arial-BoldMT  and text == 2024, or 2025
# 2, table contents and commentary has same 9.0, TimesNewRomanPSMT, do differ them and print out commentary and not table concent (table already done)
# 3. if text == "•" or text == "●" or  text == "–"

# 4. for 2019 MVR rule, catalogue front size and color are still in black color, choose this black catalogue and print them out
# 5, for figures, fix them and print out all figures (most cases only have one figure with red figure text)
# 6, "CambriaMath":     for small font size formula, in vertical direction
                    #y_middle = (y0+y1)/2
                    #y0_new = y_middle -0.6*(y1-y_middle)
                    #y1_new =  y_middle +0.6*(y1-y_middle)
# 7, for this 2024 and 2025 rule, figure extraction changed back to [:-1], the last figure is page outline
# 8, run vs code under administrator role, sometime need to split  large pdf (>1000 papges) into several smaller parts


import io
import os
import pymupdf
import json
import matplotlib
import numpy as np
import fitz  
from docx import Document
from PIL import Image
from io import BytesIO 
from docx.shared import Inches
import cv2
import pdfplumber
matplotlib.use('TkAgg',force=True)

# Methodology
# 1. use pymupdf to detect every red content, filter some noise and fake dot red content
# 2. from the font size, bold and name know where the rec content should be, extract level 1, 2, 3 flag , and figure information and check they are red colored or not, save it if not red
# 3. when encounting every red print out content, attach level 1 , 2, 3 and figure if necssary ( if they not covered by red content)
# 4. firstly print out all extracted red contents in a temporary pdf, finally stack all these net pdf contents into a word file


# 1a. load pdf file and extract all figures

# Specify the relative path to your PDF in the subfolder
pdf_path = os.path.join('doc', '1-mvr-part-4-jan25_chapter89.pdf')   ######_July_2019_split_1_50  MVR_Part_3_July_2019_527_541

doc = fitz.open(pdf_path)

#firstly extract all figures and save in a figure folder
figure_dpi = {}
dpi_resol = 380

def deblur_image(image_bytes):
    # Convert bytes to OpenCV image (numpy array)
    image = Image.open(BytesIO(image_bytes))
    image = np.array(image)
    
    # Example kernel for sharpening/deblurring
    kernel = np.array([[0, -1, 0], [-1, 5,-1], [0, -1, 0]])  # Simple sharpening kernel
    # Apply the kernel to the image
    deblurred_image = cv2.filter2D(image, -1, kernel)
    return deblurred_image

font_size_flag = False



# 2, loop page by page , detect all red contents
new_pdf = fitz.open()

for page_num in range(doc.page_count):  #doc.page_count
    page_has_redtable = False
    tableposted_flag = False
    new_page =new_pdf.new_page(width = doc[page_num].rect.width, height =doc[page_num].rect.height)
    page_flag = False
    text_blocks = doc[page_num].get_text("dict", flags=pymupdf.TEXTFLAGS_TEXT)["blocks"]
    flag_2 = False   # label level 2 outline is black or red
    flag_3 = False   # label level 3 outline is black or red
    page_figure_shown = False

    

    # extract red content
    for i, block in enumerate(text_blocks):
        x0, y0, x1, y1, = block['bbox']
        rect = fitz.Rect(x0, y0, x1, y1)
        level_font_dict = set()
        block_flag = False
        flag_2_red_done = False 
        flag_3_red_done = False
        figure_flag = False
        figure_list = []
        block_printed =  False


        for line in block["lines"]:
            for span in line["spans"]:
                text = span["text"]
                font_size = span.get("size", None)
                font_name = span.get("font")
                color = pymupdf.sRGB_to_rgb(span["color"])
                is_bold = "Bold" in font_name or"Black" in font_name
                level_font_dict.add((font_size, font_name))
                # red color text or not
                if color == (218, 31, 51):
                    if text == "•" or text == "●" or  text == "∙"or  text == "–": 
                        continue
                    else:
                        if not page_flag:
                            print(f"page numer: {page_num+1}")
                            page_flag = True
                        block_flag = True
                        if ((font_size == 9 and font_name == "TimesNewRomanPSMT") or (font_size == 12 and font_name == "Arial-ItalicMT" and (text == "2024" or text == "2025"))) and (not tableposted_flag):    ###
                            with pdfplumber.open(pdf_path) as pdf:
                                for page_number, page in enumerate(pdf.pages, start=1):
                                    if page_number == page_num +1:
                                        tables  = page.find_tables()
                                        if not tables:
                                            continue
                                        page_has_redtable = True
                                        for table_index, table in enumerate(tables):
                                            # Get the bounding box of the table
                                            x0, top, x1, bottom = table.bbox

                                            # Crop the table region from the page
                                            pix = doc[page_num].get_pixmap( dpi = 380, clip=(x0, top, x1, bottom))
                                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                                            dpi_original = img.info.get("dpi", (72, 72))
                                            print(dpi_original)

                                            rect_data = {
                                                        "x0": x0,
                                                        "y0": top,
                                                        "x1": x1,
                                                        "y1": bottom
                                                        }
                                            
                                            #new_page.insert_image(fitz.Rect(rect_data), pixmap = pix)
                                            new_page.insert_image(fitz.Rect(x0, top, x1, bottom ), pixmap = pix)

                                        tableposted_flag = True
                elif color == (0, 0 ,0) and not block_printed:
                    if is_bold == True and ((font_size == 14 and font_name == "Arial-Black") or (font_size == 36 and font_name == "Arial-Black")):
                        image = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect))
                        new_page.insert_image(fitz.Rect(rect), pixmap = image)
                        block_printed = True





        if block_flag:  # detect this red content itself is level one or level two outline or not
            if (11.0, "Arial-BoldMT") in level_font_dict: 
                flag_2_red_done = True   # level 2 outline labled by red or not
            elif (10.0, "Arial-BoldMT") in level_font_dict: 
                flag_3_red_done = True


        ### write down all font info from this block, extract level 2 and level 3 stuff, if if already labelled by red, then start from 0
        level_font_dict_all = set()
    
        for line in block["lines"]:
            for span in line["spans"]:
                text = span["text"]
                font_size = span.get("size", None)
                font_name = span.get("font")
                color = pymupdf.sRGB_to_rgb(span["color"])
                is_bold = "Bold" in font_name or "Black" in font_name
                level_font_dict_all.add((font_size, font_name))

                if font_size == 10:
                    font_size_flag = True

        if block_flag:
            if (11.0, "Arial-BoldMT") in level_font_dict_all:   
                flag_2 , flag_3 = False , False
                font_size_flag = False

            if (10.0, "Arial-BoldMT") in level_font_dict_all:  
                flag_3 = False 
            if (12.0, "Arial-BoldMT") in level_font_dict_all:  # detect is figure related text or not, if is, the insert all figures of this page here,normally only one
                figure_flag = True


        #if goes to another chapter, then all level 2 and level 3 flag == 0, form the beginning set up 0
        if (12.0, "Arial-Black") in level_font_dict_all:
            flag_2 , flag_3 = False , False
        else:       # save the level 2 and level 3 outline if not in red
            if (11.0, "Arial-BoldMT") in level_font_dict_all and not flag_2_red_done:

                rect_2 = rect
                image_2 = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect_2))
                image_2.save("output_image2.png")
                rect_data = {
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1
                    }
                with open('rect_for_crop_2.json', 'w') as f:
                    json.dump(rect_data, f)
                flag_2 = True

            if (10.0, "Arial-BoldMT") in level_font_dict_all and not flag_3_red_done:
                
                rect_3 = rect
                image_3 = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect_3))
                image_3.save("output_image3.png")
                rect_data = {
                    "x0": rect.x0,
                    "y0": rect.y0,
                    "x1": rect.x1,
                    "y1": rect.y1
                    } 
                with open('rect_for_crop_3.json', 'w') as f:
                    json.dump(rect_data, f)
                #new_page.draw_rect(rect_3, color = (1, 0, 0), width =1)
                flag_3 = True

        # finally decide what exactly need to print out for this red content area in total
        if block_flag:
            if not flag_2_red_done and flag_2 and os.path.exists('rect_for_crop_2.json'): # if level 2 not red color, print out it
                with open('rect_for_crop_2.json', 'r') as f:
                    loaded_rect_data = json.load(f)  
                rect_2 = fitz.Rect(loaded_rect_data['x0'], loaded_rect_data['y0'], loaded_rect_data['x1'], loaded_rect_data['y1'])    
                new_page.insert_image(fitz.Rect(rect_2), pixmap = fitz.Pixmap("output_image2.png"))
                os.remove('rect_for_crop_2.json')
                
            if not flag_3_red_done and flag_3 and os.path.exists('rect_for_crop_3.json'): # if level 3 not red color, print out it
                with open('rect_for_crop_3.json', 'r') as f:
                    loaded_rect_data = json.load(f)  
                rect_3 = fitz.Rect(loaded_rect_data['x0'], loaded_rect_data['y0'], loaded_rect_data['x1'], loaded_rect_data['y1'])
                new_page.insert_image(fitz.Rect(rect_3), pixmap = fitz.Pixmap("output_image3.png"))
                os.remove('rect_for_crop_3.json')

            if figure_flag: # if figure title is red color , print out this figure
                if not page_figure_shown:
                    image = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect))
                    new_page.insert_image(fitz.Rect(rect), pixmap = image)

                    print("has red figure here")
                    image_list = doc[page_num].get_images(full=True)
                    print(len(image_list))


                    #if len(image_list) <= 1: # every pafge has a page format figure which is not a real figure
                        #continue
                    for img_index, img in enumerate(image_list[:-1]):
        
                        xref = img[0]  # Get image reference
                        base_image = doc.extract_image(xref)  # Extract image as a dictionary
                        # Get image bytes and convert to PIL Image
                        image_bytes = base_image["image"]
                        image = Image.open(io.BytesIO(image_bytes))
                        width, height = image.size
                        dpi = image.info.get("dpi", (96.012, 96.012))

                        rect = doc[page_num].get_image_bbox(img)   ## save the image and locatio in this page
                        #new_page.insert_image(fitz.Rect(rect), pixmap = fitz.Pixmap(image))   

                        new_page.insert_image(fitz.Rect(rect), pixmap = doc[page_num].get_pixmap( dpi = 380, clip=rect))
                        page_figure_shown = True
            else:

                #  dont'r print out the text inside tables, otherwise print all red text block out
                if page_has_redtable and ((9.0, "TimesNewRomanPSMT") in level_font_dict_all):  
                    continue
                else:
                    if ((8.0, "CambriaMath") in level_font_dict_all) or ((10.5, "CambriaMath") in level_font_dict_all):
                        y_middle = (y0+y1)/2
                        y0_new = y_middle -0.36*(y1-y_middle)
                        y1_new =  y_middle +0.36*(y1-y_middle)
                        rect  = fitz.Rect(x0, y0_new, x1, y1_new)
                       
                    image = doc[page_num].get_pixmap(dpi = dpi_resol, clip=fitz.Rect(rect))
                    new_page.insert_image(fitz.Rect(rect), pixmap = image)
            
    new_pdf_path = "doc/New_red_contents_v3_{}.pdf".format(pdf_path[4:].split(".")[0])  
    # if os.path.exists(new_pdf_path): 
    #     os.remove(new_pdf_path)
    new_pdf.save(new_pdf_path)


# 3. scan the generated pdf intermediate file and stack the contents into word format
#  firstly scan all as images, need to scale down for some large images which is too large for the word page

def extract_images_with_positions(pdf_path):
    # Open the PDF file
    pdf_document = fitz.open(pdf_path)
    images = []

    # Iterate through each page
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        image_list = page.get_images(full=True)

        # Iterate through each image on the page
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]

            # Get the position of the image on the page
            image_rects = page.get_image_rects(xref)
            for rect in image_rects:
                images.append({
                    "image_bytes": image_bytes,
                    "image_ext": image_ext,
                    "x0": rect.x0,  # Left position
                    "y0": rect.y0,  # Top position
                    "x1": rect.x1,  # Right position
                    "y1": rect.y1,  # Bottom position
                    "width": rect.width,
                    "height": rect.height
                })
    return images

def create_word_document_with_positioned_images(images, output_docx_path):
    # Create a new Word document
    doc = Document()

    for image_info in images:
        image_bytes = image_info["image_bytes"]
        x0 = image_info["x0"]
        width = image_info["width"]
        height = image_info["height"]

        # Add a paragraph for the image
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # Add the image to the paragraph
        image_stream = io.BytesIO(image_bytes)
        run.add_picture(image_stream, width=Inches(width / 86))  # Convert points to inches

        # Set the horizontal position of the image
        paragraph.alignment = 0  # Left alignment
        paragraph.paragraph_format.left_indent = Inches(x0 / 86)  # Set left indent based on x0

    # Save the Word document
    doc.save(output_docx_path)

def main(pdf_path, output_docx_path):
    # Extract images with their positions from the PDF
    images = extract_images_with_positions(pdf_path)

    # Create a Word document with the extracted images at their original horizontal positions
    create_word_document_with_positioned_images(images, output_docx_path)

if __name__ == "__main__":
    word_file_name = "doc/word_output_{}.docx".format(pdf_path[4:].split(".")[0])
    main(new_pdf_path, word_file_name)



